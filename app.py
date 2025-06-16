import streamlit as st
from sentence_transformers import SentenceTransformer, util
import docx
from docx import Document
import difflib
import numpy as np
import tempfile
import os

st.title("📄 Legal Clause Comparison Tool")
st.write("Upload two versions of a Word (.docx) document to compare clauses and summarize changes.")

# Upload files
file_a = st.file_uploader("Upload Version A", type=["docx"])
file_b = st.file_uploader("Upload Version B", type=["docx"])

def extract_clauses_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

def summarize_diff(text_a, text_b):
    if text_a is None:
        return "Clause added in version B."
    elif text_b is None:
        return "Clause removed in version B."
    elif text_a == text_b:
        return "No change."
    else:
        diff = difflib.ndiff(text_a.split(), text_b.split())
        changes = [d for d in diff if d.startswith('+ ') or d.startswith('- ')]
        return "Changes: " + ' '.join(changes)

if file_a and file_b:
    with st.spinner("Processing and comparing clauses..."):
        # Load clauses
        clauses_a = extract_clauses_from_docx(file_a)
        clauses_b = extract_clauses_from_docx(file_b)

        # Encode with SentenceTransformer
        model = SentenceTransformer("all-MiniLM-L6-v2")
        embeddings_a = model.encode(clauses_a, convert_to_tensor=True)
        embeddings_b = model.encode(clauses_b, convert_to_tensor=True)

        # Match clauses
        matched_pairs = []
        used_b = set()

        for idx_a, emb_a in enumerate(embeddings_a):
            similarities = util.cos_sim(emb_a, embeddings_b)[0].cpu().numpy()
            best_match_idx = int(np.argmax(similarities))
            best_score = similarities[best_match_idx]

            if best_score > 0.75 and best_match_idx not in used_b:
                matched_pairs.append({
                    "unified_id": len(matched_pairs) + 1,
                    "text_a": clauses_a[idx_a],
                    "text_b": clauses_b[best_match_idx],
                    "match_score": best_score
                })
                used_b.add(best_match_idx)
            else:
                matched_pairs.append({
                    "unified_id": len(matched_pairs) + 1,
                    "text_a": clauses_a[idx_a],
                    "text_b": None,
                    "match_score": best_score
                })

        for idx_b, clause_b in enumerate(clauses_b):
            if idx_b not in used_b:
                matched_pairs.append({
                    "unified_id": len(matched_pairs) + 1,
                    "text_a": None,
                    "text_b": clause_b,
                    "match_score": None
                })

        # Summarize
        for pair in matched_pairs:
            pair["summary"] = summarize_diff(pair["text_a"], pair["text_b"])

        # Create docx report
        output_doc = Document()
        output_doc.add_heading("Clause Comparison Report", level=1)

        table = output_doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Version A'
        hdr_cells[1].text = 'Version B'
        hdr_cells[2].text = 'Summary of Changes'

        for pair in matched_pairs:
            row_cells = table.add_row().cells
            row_cells[0].text = pair['text_a'] if pair['text_a'] else "(Not present)"
            row_cells[1].text = pair['text_b'] if pair['text_b'] else "(Not present)"
            row_cells[2].text = pair['summary']

        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            output_doc.save(tmp.name)
            tmp_path = tmp.name

    # Offer download
    with open(tmp_path, "rb") as f:
        st.success("✅ Comparison complete!")
        st.download_button(
            label="📥 Download Clause Comparison Report",
            data=f,
            file_name="clause_comparison_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

